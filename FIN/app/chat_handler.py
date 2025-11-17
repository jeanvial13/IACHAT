import os
import time
import io
from datetime import datetime, timedelta
from flask import (
    Flask,
    request,
    jsonify,
    render_template,
    send_file,
    redirect,
    url_for,
    session,
)
from werkzeug.utils import secure_filename
from openai import OpenAI

try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None

try:
    from docx import Document
except Exception:
    Document = None

try:
    from openpyxl import Workbook
except Exception:
    Workbook = None

try:
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.pagesizes import A4
except Exception:
    SimpleDocTemplate = None

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

app = Flask(__name__, static_folder="static", template_folder="templates")

UPLOAD_FOLDER = os.path.join(BASE_DIR, "uploads")
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER

app.secret_key = os.environ.get("FLASK_SECRET_KEY", "dev-secret")

OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY")
DEFAULT_MODEL = os.environ.get("OPENAI_MODEL", "gpt-4o-mini")

client = OpenAI(api_key=OPENAI_API_KEY)

LOG_FILE = os.path.join(BASE_DIR, "server.log")
DEMS_FILE = os.path.join(BASE_DIR, "dem_projects.json")


# ---------------- Utilities ----------------


def _log(line: str) -> None:
    ts = datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S")
    try:
        with open(LOG_FILE, "a", encoding="utf-8") as f:
            f.write(f"[{ts}] {line}\n")
    except Exception:
        pass


def extract_text(path: str) -> str:
    """Extrae texto de TXT, PDF, DOCX y DOC."""
    lower = path.lower()
    try:
        # TXT / MD
        if lower.endswith(".txt") or lower.endswith(".md"):
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()

        # PDF
        if lower.endswith(".pdf") and PdfReader is not None:
            reader = PdfReader(path)
            parts = []
            for page in reader.pages:
                try:
                    parts.append(page.extract_text() or "")
                except:
                    pass
            return "\n".join(parts)

        # DOCX
        if lower.endswith(".docx") and Document is not None:
            doc = Document(path)
            return "\n".join(p.text for p in doc.paragraphs)

        # DOC (legacy)
        if lower.endswith(".doc"):
            try:
                import textract
                return textract.process(path).decode("utf-8", errors="ignore")
            except:
                return ""
    except Exception as e:
        _log(f"Error leyendo archivo {path}: {e}")

    return ""



def load_dems():
    if not os.path.exists(DEMS_FILE):
        return []
    try:
        with open(DEMS_FILE, "r", encoding="utf-8") as f:
            data = f.read().strip()
            if not data:
                return []
            return json.loads(data)
    except Exception as e:
        _log(f"Error loading DEM file: {e}")
        return []


def save_dems(dems):
    try:
        with open(DEMS_FILE, "w", encoding="utf-8") as f:
            json.dump(dems, f, ensure_ascii=False, indent=2)
    except Exception as e:
        _log(f"Error saving DEM file: {e}")


def enrich_dem(dem):
    dem = dict(dem)
    start = dem.get("start_date")
    dem["duration_days"] = None
    if start:
        try:
            dt = datetime.strptime(start, "%Y-%m-%d")
            dem["duration_days"] = (datetime.utcnow().date() - dt.date()).days
        except Exception:
            pass

    notes = dem.get("notes") or []
    dem["last_note"] = notes[-1] if notes else ""

    updated_str = dem.get("updated_at") or dem.get("created_at")
    sla_breached = False
    if updated_str:
        try:
            upd = datetime.fromisoformat(updated_str)
            if datetime.utcnow() - upd > timedelta(days=5):
                sla_breached = True
        except Exception:
            pass
    dem["sla_breached"] = sla_breached
    if "archived" not in dem:
        dem["archived"] = False
    return dem


def build_portfolio_text(dems):
    if not dems:
        return "No hay proyectos DEM registrados actualmente."

    lines = []
    lines.append("DEMS Portfolio Report")
    lines.append("")
    lines.append(
        "Este documento resume el estado actual de los proyectos DEM activos y archivados, "
        "incluyendo su situación general, Workflow Status vigente y el responsable de la siguiente actividad."
    )
    lines.append("")
    for dem in dems:
        e = enrich_dem(dem)
        lines.append(f"DEM: {e.get('name','(sin nombre)')}")
        lines.append(f"  Project Title: {e.get('title','')}")
        lines.append(f"  Sponsor: {e.get('sponsor','-')}  |  Requester: {e.get('requester','-')}")
        lines.append(f"  BA Owner: {e.get('ba_owner','-')}  |  Current Task Owner: {e.get('current_owner','-')}")
        lines.append(f"  Cost Center: {e.get('cost_center','-')}")
        lines.append(f"  Start Date: {e.get('start_date','-')}  |  Duration (days): {e.get('duration_days')}")
        lines.append(f"  DEM Status: {e.get('status','-')}")
        lines.append(f"  Workflow Status: {e.get('workflow_status','-')}")
        lines.append(
            "  SLA: " + ("SLA Breached" if e.get("sla_breached") else "SLA OK")
        )
        notes = e.get("notes") or []
        if notes:
            last_two = notes[-2:]
            lines.append("  Last Notes:")
            for n in last_two:
                lines.append(f"    - {n}")
        else:
            lines.append("  Last Notes:  (no notes registered)")
        lines.append("")
    return "\n".join(lines)


# ---------------- Auth & Views ----------------

import json


@app.route("/login", methods=["GET", "POST"])
def login():
    error = None
    if request.method == "POST":
        username = (request.form.get("username") or "").strip()
        password = request.form.get("password") or ""
        app_user = os.environ.get("APP_USER")
        app_pass = os.environ.get("APP_PASS")

        if app_user and app_pass and username == app_user and password == app_pass:
            session["auth"] = True
            return redirect(url_for("home"))
        error = "Usuario o contraseña incorrectos."
    return render_template("login.html", error=error)


@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))


def require_auth():
    if not session.get("auth"):
        return redirect(url_for("login"))
    return None


@app.route("/")
def home():
    maybe = require_auth()
    if maybe is not None:
        return maybe
    return render_template("index.html")


@app.route("/dems")
def dems_page():
    maybe = require_auth()
    if maybe is not None:
        return maybe
    return render_template("dem_manager.html")


# ---------------- Chat & Files ----------------


@app.route("/chat", methods=["POST"])
def chat():
    maybe = require_auth()
    if maybe is not None:
        return jsonify({"error": "No autorizado"}), 401

    data = request.get_json() or {}
    message = data.get("message", "").strip()
    history = data.get("history", [])
    model = data.get("model") or DEFAULT_MODEL
    file_summaries = data.get("file_summaries", [])

    if not message:
        return jsonify({"error": "Mensaje vacío"}), 400

    messages = []
    for m in history:
        role = m.get("role")
        content = m.get("content")
        if role in ("user", "assistant") and isinstance(content, str):
            messages.append({"role": role, "content": content})

    user_content = message

    if file_summaries:
        user_content += "\n\n[Contexto de archivos adjuntos]\n"
        for fsum in file_summaries:
            fname = fsum.get("filename", "archivo")
            summ = fsum.get("summary", "")
            user_content += f"- {fname}: {summ}\n"

    messages.append({"role": "user", "content": user_content})

    try:
        completion = client.chat.completions.create(
            model=model,
            messages=messages,
        )
        answer = completion.choices[0].message.content
        return jsonify({"reply": answer})
    except Exception as e:
        _log(f"Error en OpenAI chat: {e}")
        return jsonify({"error": "Error al llamar al modelo de IA."}), 500


@app.route("/upload", methods=["POST"])
def upload():
    maybe = require_auth()
    if maybe is not None:
        return jsonify({"error": "No autorizado"}), 401

    if "files" not in request.files:
        return jsonify({"error": "No se enviaron archivos"}), 400

    files = request.files.getlist("files")
    if not files:
        return jsonify({"error": "No se enviaron archivos"}), 400

    results = []
    for f in files:
        filename = secure_filename(f.filename or "archivo")
        save_name = f"{int(time.time())}_{filename}"
        path = os.path.join(app.config["UPLOAD_FOLDER"], save_name)
        _log(f"Guardando archivo en {path}")

        f.save(path)

        text = extract_text(path)
        if not text:
            results.append(
                {
                    "filename": filename,
                    "summary": "No pude leer este archivo (formato no soportado o vacío).",
                }
            )
            continue

        try:
            completion = client.chat.completions.create(
                model=DEFAULT_MODEL,
                messages=[
                    {
                        "role": "user",
                        "content": (
                            "Resume el siguiente documento en pocas frases, "
                            "resaltando puntos clave para trabajo de análisis y gestión de proyectos:\n\n"
                            f"{text[:8000]}"
                        ),
                    }
                ],
            )
            summary = completion.choices[0].message.content
        except Exception as e:
            _log(f"Error resumiendo archivo: {e}")
            summary = "No se pudo generar resumen automático, pero el archivo se recibió correctamente."

        results.append({"filename": filename, "summary": summary})

    return jsonify({"files": results})


# ---------------- DEM / Project Manager API ----------------


def get_dems_filtered(archived: bool):
    dems = load_dems()
    return [enrich_dem(d) for d in dems if bool(d.get("archived", False)) == archived]


@app.route("/api/dems/projects", methods=["GET"])
def list_dems():
    maybe = require_auth()
    if maybe is not None:
        return jsonify({"error": "No autorizado"}), 401

    archived_str = request.args.get("archived", "false").lower()
    archived = archived_str in ("1", "true", "yes")
    projects = get_dems_filtered(archived)
    return jsonify({"projects": projects})


@app.route("/api/dems/projects", methods=["POST"])
def create_dem():
    maybe = require_auth()
    if maybe is not None:
        return jsonify({"error": "No autorizado"}), 401

    data = request.get_json() or {}
    now_iso = datetime.utcnow().isoformat()

    dem = {
        "id": f"dem_{int(time.time() * 1000)}",
        "name": data.get("name", "").strip(),
        "sponsor": data.get("sponsor", "").strip(),
        "requester": data.get("requester", "").strip(),
        "ba_owner": data.get("ba_owner", "").strip(),
        "title": data.get("title", "").strip(),
        "change_request": data.get("change_request", "").strip(),
        "cost_center": data.get("cost_center", "").strip(),
        "status": data.get("status", "").strip() or "Idea",
        "workflow_status": data.get("workflow_status", "").strip() or "Intake",
        "current_owner": data.get("current_owner", "").strip(),
        "start_date": data.get("start_date", "").strip(),
        "notes": [],
        "doc_summary": "",
        "created_at": now_iso,
        "updated_at": now_iso,
        "archived": False,
    }

    initial_note = data.get("initial_note", "").strip() if "initial_note" in data else ""
    if initial_note:
        dem["notes"].append(initial_note)

    dems = load_dems()
    dems.append(dem)
    save_dems(dems)

    return jsonify({"project": enrich_dem(dem)})


def _update_dem(id, updater):
    dems = load_dems()
    for i, d in enumerate(dems):
        if d.get("id") == id:
            updater(d)
            d["updated_at"] = datetime.utcnow().isoformat()
            dems[i] = d
            save_dems(dems)
            return enrich_dem(d)
    return None


@app.route("/api/dems/projects/<id>/note", methods=["POST"])
def add_dem_note(id):
    maybe = require_auth()
    if maybe is not None:
        return jsonify({"error": "No autorizado"}), 401

    data = request.get_json() or {}
    text = (data.get("text") or "").strip()
    if not text:
        return jsonify({"error": "Nota vacía."}), 400

    def updater(d):
        notes = d.get("notes") or []
        notes.append(text)
        d["notes"] = notes

    project = _update_dem(id, updater)
    if not project:
        return jsonify({"error": "DEM no encontrado."}), 404
    return jsonify({"project": project})


@app.route("/api/dems/projects/<id>/update", methods=["POST"])
def update_dem(id):
    maybe = require_auth()
    if maybe is not None:
        return jsonify({"error": "No autorizado"}), 401

    data = request.get_json() or {}

    def updater(d):
        for field in [
            "name",
            "sponsor",
            "requester",
            "ba_owner",
            "title",
            "change_request",
            "cost_center",
            "status",
            "workflow_status",
            "current_owner",
            "start_date",
        ]:
            if field in data:
                d[field] = (data.get(field) or "").strip()

    project = _update_dem(id, updater)
    if not project:
        return jsonify({"error": "DEM no encontrado."}), 404
    return jsonify({"project": project})


@app.route("/api/dems/projects/<id>/archive", methods=["POST"])
def archive_dem(id):
    maybe = require_auth()
    if maybe is not None:
        return jsonify({"error": "No autorizado"}), 401

    def updater(d):
        d["archived"] = True

    project = _update_dem(id, updater)
    if not project:
        return jsonify({"error": "DEM no encontrado."}), 404
    return jsonify({"project": project})


@app.route("/api/dems/projects/<id>/restore", methods=["POST"])
def restore_dem(id):
    maybe = require_auth()
    if maybe is not None:
        return jsonify({"error": "No autorizado"}), 401

    def updater(d):
        d["archived"] = False

    project = _update_dem(id, updater)
    if not project:
        return jsonify({"error": "DEM no encontrado."}), 404
    return jsonify({"project": project})


@app.route("/api/dems/projects/<id>/delete", methods=["POST"])
def delete_dem(id):
    maybe = require_auth()
    if maybe is not None:
        return jsonify({"error": "No autorizado"}), 401

    dems = load_dems()
    new_dems = [d for d in dems if d.get("id") != id]
    if len(new_dems) == len(dems):
        return jsonify({"error": "DEM no encontrado."}), 404
    save_dems(new_dems)
    return jsonify({"success": True})


@app.route("/api/dems/projects/<id>/attach", methods=["POST"])
def attach_doc(id):
    maybe = require_auth()
    if maybe is not None:
        return jsonify({"error": "No autorizado"}), 401

    if "file" not in request.files:
        return jsonify({"error": "No se recibió archivo."}), 400

    file = request.files["file"]
    filename = secure_filename(file.filename or "document")
    save_name = f"{int(time.time())}_{filename}"
    path = os.path.join(app.config["UPLOAD_FOLDER"], save_name)
    file.save(path)

    text = extract_text(path)
    if not text:
        return jsonify({"error": "No pude leer este archivo para generar un resumen."}), 400

    summary = ""
    try:
        completion = client.chat.completions.create(
            model=DEFAULT_MODEL,
            messages=[
                {
                    "role": "user",
                    "content": (
                        "Genera un resumen ejecutivo del siguiente contenido, "
                        "resaltando objetivo, alcance, riesgos, responsables y próximos pasos. "
                        "No menciones que eres un modelo de IA ni describas el proceso, solo entrega el resumen.\n\n"
                        f"{text[:8000]}"
                    ),
                }
            ],
        )
        summary = completion.choices[0].message.content
    except Exception as e:
        _log(f"Error generating doc summary: {e}")
        summary = "No se pudo generar un resumen automático, pero el documento quedó adjunto a este DEM."

    def updater(d):
        d["doc_summary"] = summary

    project = _update_dem(id, updater)
    if not project:
        return jsonify({"error": "DEM no encontrado."}), 404
    return jsonify({"project": project})


@app.route("/api/dems/export", methods=["GET"])
def export_active_excel():
    maybe = require_auth()
    if maybe is not None:
        return jsonify({"error": "No autorizado"}), 401

    if Workbook is None:
        return jsonify({"error": "openpyxl no está disponible en el servidor."}), 500

    wb = Workbook()
    ws = wb.active
    ws.title = "Active DEMs"

    headers = [
        "ID",
        "Name",
        "Title",
        "Sponsor",
        "Requester",
        "BA Owner",
        "Cost Center",
        "Status",
        "Workflow Status",
        "Current Task Owner",
        "Start Date",
        "Duration Days",
        "SLA",
        "Last Note",
    ]
    ws.append(headers)

    for dem in get_dems_filtered(False):
        ws.append(
            [
                dem.get("id"),
                dem.get("name"),
                dem.get("title"),
                dem.get("sponsor"),
                dem.get("requester"),
                dem.get("ba_owner"),
                dem.get("cost_center"),
                dem.get("status"),
                dem.get("workflow_status"),
                dem.get("current_owner"),
                dem.get("start_date"),
                dem.get("duration_days"),
                "Breached" if dem.get("sla_breached") else "OK",
                dem.get("last_note"),
            ]
        )

    bio = io.BytesIO()
    wb.save(bio)
    bio.seek(0)
    return send_file(
        bio,
        as_attachment=True,
        download_name="dems_active.xlsx",
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


@app.route("/api/dems/export_archived", methods=["GET"])
def export_archived_excel():
    maybe = require_auth()
    if maybe is not None:
        return jsonify({"error": "No autorizado"}), 401

    if Workbook is None:
        return jsonify({"error": "openpyxl no está disponible en el servidor."}), 500

    wb = Workbook()
    ws = wb.active
    ws.title = "Archived DEMs"

    headers = [
        "ID",
        "Name",
        "Title",
        "Sponsor",
        "Requester",
        "BA Owner",
        "Cost Center",
        "Status",
        "Workflow Status",
        "Current Task Owner",
        "Start Date",
        "Duration Days",
        "SLA",
        "Last Note",
    ]
    ws.append(headers)

    for dem in get_dems_filtered(True):
        ws.append(
            [
                dem.get("id"),
                dem.get("name"),
                dem.get("title"),
                dem.get("sponsor"),
                dem.get("requester"),
                dem.get("ba_owner"),
                dem.get("cost_center"),
                dem.get("status"),
                dem.get("workflow_status"),
                dem.get("current_owner"),
                dem.get("start_date"),
                dem.get("duration_days"),
                "Breached" if dem.get("sla_breached") else "OK",
                dem.get("last_note"),
            ]
        )

    bio = io.BytesIO()
    wb.save(bio)
    bio.seek(0)
    return send_file(
        bio,
        as_attachment=True,
        download_name="dems_archived.xlsx",
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


@app.route("/api/dems/report", methods=["POST"])
def dem_report():
    maybe = require_auth()
    if maybe is not None:
        return jsonify({"error": "No autorizado"}), 401

    dems = load_dems()
    text = build_portfolio_text(dems)
    return jsonify({"report": text})


@app.route("/api/dems/download/<fmt>", methods=["GET"])
def dem_download(fmt):
    maybe = require_auth()
    if maybe is not None:
        return jsonify({"error": "No autorizado"}), 401

    fmt = fmt.lower()
    if fmt not in ("txt", "pdf", "docx"):
        return jsonify({"error": "Formato no soportado."}), 400

    dems = load_dems()
    text = build_portfolio_text(dems)

    if fmt == "txt":
        bio = io.BytesIO()
        bio.write(text.encode("utf-8"))
        bio.seek(0)
        return send_file(
            bio,
            as_attachment=True,
            download_name="dems_portfolio.txt",
            mimetype="text/plain; charset=utf-8",
        )

    if fmt == "docx":
        if Document is None:
            return jsonify({"error": "python-docx no está disponible."}), 500
        doc = Document()
        doc.add_heading("DEMS Portfolio Report", level=1)
        doc.add_paragraph(
            "Resumen del estado actual de los proyectos DEM, incluyendo situación general, Workflow Status y Current Task Owner."
        )
        doc.add_paragraph("")
        for line in text.split("\n"):
            doc.add_paragraph(line)

        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        return send_file(
            bio,
            as_attachment=True,
            download_name="dems_portfolio.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    if fmt == "pdf":
        if SimpleDocTemplate is None:
            return jsonify({"error": "reportlab no está disponible."}), 500

        bio = io.BytesIO()
        doc = SimpleDocTemplate(bio, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        story.append(Paragraph("DEMS Portfolio Report", styles["Title"]))
        story.append(Spacer(1, 12))
        for paragraph in text.split("\n\n"):
            story.append(Paragraph(paragraph.replace("\n", "<br />"), styles["Normal"]))
            story.append(Spacer(1, 8))
        doc.build(story)
        bio.seek(0)
        return send_file(
            bio,
            as_attachment=True,
            download_name="dems_portfolio.pdf",
            mimetype="application/pdf",
        )

    return jsonify({"error": "Formato no soportado."}), 400


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=8080)
