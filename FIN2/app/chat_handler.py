import os
import time
import io
import json
from datetime import datetime, timedelta
from flask import (
    Flask,
    request,
    jsonify,
    render_template,
    send_file,
    redirect,
    url_for,
    session,
)
from werkzeug.utils import secure_filename
from openai import OpenAI

try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None

try:
    from docx import Document
except Exception:
    Document = None

try:
    from openpyxl import Workbook
except Exception:
    Workbook = None

try:
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.pagesizes import A4
except Exception:
    SimpleDocTemplate = None

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

app = Flask(__name__, static_folder="static", template_folder="templates")

UPLOAD_FOLDER = os.path.join(BASE_DIR, "uploads")
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER

app.secret_key = os.environ.get("FLASK_SECRET_KEY", "dev-secret")

OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY")
DEFAULT_MODEL = os.environ.get("OPENAI_MODEL", "gpt-4o-mini")

client = OpenAI(api_key=OPENAI_API_KEY)

LOG_FILE = os.path.join(BASE_DIR, "server.log")
DEMS_FILE = os.path.join(BASE_DIR, "dem_projects.json")


# ---------------- Utilities ----------------


def _log(line: str) -> None:
    ts = datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S")
    try:
        with open(LOG_FILE, "a", encoding="utf-8") as f:
            f.write(f"[{ts}] {line}\n")
    except Exception:
        pass


def extract_text(path: str) -> str:
    """Extract text from TXT, PDF, DOCX and DOC."""
    lower = path.lower()
    try:
        # TXT / MD
        if lower.endswith(".txt") or lower.endswith(".md"):
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()

        # PDF
        if lower.endswith(".pdf") and PdfReader is not None:
            reader = PdfReader(path)
            parts = []
            for page in reader.pages:
                try:
                    parts.append(page.extract_text() or "")
                except Exception:
                    pass
            return "\n".join(parts)

        # DOCX
        if lower.endswith(".docx") and Document is not None:
            doc = Document(path)
            return "\n".join(p.text for p in doc.paragraphs)

        # DOC (legacy)
        if lower.endswith(".doc"):
            try:
                import textract

                return textract.process(path).decode("utf-8", errors="ignore")
            except Exception:
                return ""
    except Exception as e:
        _log(f"Error reading file {path}: {e}")

    return ""


def load_dems():
    if not os.path.exists(DEMS_FILE):
        return []
    try:
        with open(DEMS_FILE, "r", encoding="utf-8") as f:
            data = f.read().strip()
            if not data:
                return []
            return json.loads(data)
    except Exception as e:
        _log(f"Error loading DEM file: {e}")
        return []


def save_dems(dems):
    try:
        with open(DEMS_FILE, "w", encoding="utf-8") as f:
            json.dump(dems, f, ensure_ascii=False, indent=2)
    except Exception as e:
        _log(f"Error saving DEM file: {e}")


def _clean_note_text(raw: str) -> str:
    """
    Quita una fecha duplicada al inicio si ya viene en el texto.

    Ejemplo:
        "[2025-11-18 00:52] — texto" -> "texto"
    """
    if not raw:
        return ""
    txt = raw.strip()
    if txt.startswith("[") and "] — " in txt:
        end = txt.find("] — ")
        if end != -1:
            return txt[end + 4 :].lstrip()
    return txt


def _format_note(note) -> str:
    """Devuelve una nota legible con la fecha, sin duplicarla."""
    if isinstance(note, dict):
        text = _clean_note_text(note.get("text", ""))
        date = note.get("date")
        if date:
            return f"[{date}] — {text}" if text else f"[{date}]"
        return text
    # compatibilidad con notas antiguas tipo string
    return _clean_note_text(str(note))


def enrich_dem(dem):
    """
    Agrega campos calculados:
    - duration_days
    - last_note
    - sla_breached
    - archived
    - priority
    - documents (lista)
    Además limpia el texto de las notas para quitar fechas duplicadas.
    """
    dem = dict(dem)

    # Normalizar notas (sin modificar el JSON en disco)
    raw_notes = dem.get("notes") or []
    cleaned_notes = []
    for n in raw_notes:
        if isinstance(n, dict):
            nn = dict(n)
            nn["text"] = _clean_note_text(nn.get("text", ""))
            cleaned_notes.append(nn)
        else:
            cleaned_notes.append(_clean_note_text(str(n)))
    dem["notes"] = cleaned_notes
    notes = cleaned_notes

    # Duración en días
    start_date = dem.get("start_date")
    dem["duration_days"] = None
    if start_date:
        try:
            dt = datetime.strptime(start_date, "%Y-%m-%d")
            dem["duration_days"] = (datetime.utcnow().date() - dt.date()).days
        except Exception:
            pass

    # Última nota formateada
    if notes:
        dem["last_note"] = _format_note(notes[-1])
    else:
        dem["last_note"] = ""

    # SLA (5 días sin actualización)
    updated_str = dem.get("updated_at") or dem.get("created_at")
    sla_breached = False
    if updated_str:
        try:
            upd = datetime.fromisoformat(updated_str)
            if datetime.utcnow() - upd > timedelta(days=5):
                sla_breached = True
        except Exception:
            pass
    dem["sla_breached"] = sla_breached

    # Archivado
    if "archived" not in dem:
        dem["archived"] = False

    # Prioridad por defecto
    if not dem.get("priority"):
        dem["priority"] = "2"

    # Lista de documentos
    docs = dem.get("documents")
    if docs is None or not isinstance(docs, list):
        dem["documents"] = []

    return dem


def build_portfolio_text(dems):
    """
    Construye el texto corporativo del portafolio para TXT/DOCX/PDF y el panel de UI.

    Sección 1: Projects Resume (resumen ejecutivo)
    Sección 2: Projects Details (detalle por DEM)

    IMPORTANTE:
      - Solo se deben pasar DEMs activos (no archivados).
    """
    if not dems:
        return "There are currently no DEM projects registered."

    enriched = [enrich_dem(d) for d in dems]

    run_date_human = datetime.utcnow().strftime("%B %d, %Y")
    header = f"{run_date_human} — Andres Villanueva DEMS Report"

    lines = []
    lines.append(header)
    lines.append("")
    lines.append("1. Projects Resume — Executive Overview")
    lines.append("")

    total = len(enriched)

    # Distribución de prioridades
    p1 = p2 = p3 = p4 = 0
    sla_ok = sla_breached = 0
    status_counts = {}

    for e in enriched:
        pr = str(e.get("priority") or "2")
        if pr == "1":
            p1 += 1
        elif pr == "2":
            p2 += 1
        elif pr == "3":
            p3 += 1
        elif pr == "4":
            p4 += 1

        if e.get("sla_breached"):
            sla_breached += 1
        else:
            sla_ok += 1

        st = e.get("status") or "N/A"
        status_counts[st] = status_counts.get(st, 0) + 1

    # Métricas clave
    lines.append("Key portfolio metrics for active DEM projects:")
    lines.append(f"• Total active DEMs: {total}")
    lines.append("• Priority distribution:")
    lines.append(f"   – P1 (Critical): {p1}")
    lines.append(f"   – P2 (High): {p2}")
    lines.append(f"   – P3 (Medium): {p3}")
    lines.append(f"   – P4 (Low): {p4}")
    lines.append(f"• SLA window (last 5 days): OK={sla_ok} | Breached={sla_breached}")

    if status_counts:
        top_status = sorted(status_counts.items(), key=lambda kv: kv[1], reverse=True)[
            :3
        ]
        status_str = ", ".join(f"{name}: {cnt}" for name, cnt in top_status)
        lines.append(f"• Most common DEM Status: {status_str}")

    lines.append("")
    lines.append("Active DEM overview (project name + latest comment):")
    lines.append("")

    for e in enriched:
        name = e.get("name", "(no name)")
        status = e.get("status", "-")
        workflow = e.get("workflow_status", "-")
        pr = e.get("priority", "2")

        raw_notes = e.get("notes") or []
        if raw_notes:
            latest = _format_note(raw_notes[-1])
        else:
            latest = "No recent notes registered."

        lines.append(
            f"• {name} — Status: {status} | Workflow: {workflow} | Priority: P{pr}"
        )
        lines.append(f"  Last update: {latest}")
        lines.append("")

    # Línea de separación donde marcaste en rojo
    lines.append("-" * 78)
    lines.append("")
    lines.append(
        "The following pages contain a detailed section per DEM, including "
        "Project Title, Sponsor, BA Owner, Workflow Status, SLA condition and "
        "the most recent notes captured during project follow-up."
    )
    lines.append("")
    lines.append("-" * 78)
    lines.append("")
    lines.append("2. Projects Details")
    lines.append("")

    # Detalle por DEM
    for e in enriched:
        lines.append(f"DEM: {e.get('name', '(no name)')}")
        lines.append(f"Project Title: {e.get('title', '')}")
        lines.append(
            f"Sponsor: {e.get('sponsor', '-')}"
            f" | Requester: {e.get('requester', '-')}"
        )
        lines.append(
            f"BA Owner: {e.get('ba_owner', '-')}"
            f" | Current Task Owner: {e.get('current_owner', '-')}"
        )
        lines.append(f"Cost Center: {e.get('cost_center', '-')}")
        lines.append(
            f"Start Date: {e.get('start_date', '-')}"
            f" | Duration (days): {e.get('duration_days')}"
        )
        lines.append(f"DEM Status: {e.get('status', '-')}")
        lines.append(f"Workflow Status: {e.get('workflow_status', '-')}")
        lines.append(f"Priority (1–4): {e.get('priority', '2')}")
        lines.append(
            "SLA Status: "
            + (
                "SLA Breached — project requires immediate follow-up with Sponsor and IT lead."
                if e.get("sla_breached")
                else "SLA OK — project updated within acceptable window."
            )
        )

        raw_notes = e.get("notes") or []
        if raw_notes:
            formatted_notes = [_format_note(n) for n in raw_notes]
            last_two = formatted_notes[-2:]
            lines.append("Last Notes (most recent entries):")
            for n in last_two:
                lines.append(f"- {n}")
        else:
            lines.append("Last Notes: (no notes registered)")

        lines.append("")
        lines.append("-" * 78)
        lines.append("")

    return "\n".join(lines)


# ---------------- Auth & Views ----------------


@app.route("/login", methods=["GET", "POST"])
def login():
    error = None
    if request.method == "POST":
        username = (request.form.get("username") or "").strip()
        password = request.form.get("password") or ""
        app_user = os.environ.get("APP_USER")
        app_pass = os.environ.get("APP_PASS")

        if app_user and app_pass and username == app_user and password == app_pass:
            session["auth"] = True
            return redirect(url_for("home"))
        error = "Usuario o contraseña incorrectos."
    return render_template("login.html", error=error)


@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))


def require_auth():
    if not session.get("auth"):
        return redirect(url_for("login"))
    return None


@app.route("/")
def home():
    maybe = require_auth()
    if maybe is not None:
        return maybe
    return render_template("index.html")


@app.route("/dems")
def dems_page():
    maybe = require_auth()
    if maybe is not None:
        return maybe
    return render_template("dem_manager.html")


# ---------------- Chat & Files (main chat page) ----------------


@app.route("/chat", methods=["POST"])
def chat():
    maybe = require_auth()
    if maybe is not None:
        return jsonify({"error": "No autorizado"}), 401

    data = request.get_json() or {}
    message = data.get("message", "").strip()
    history = data.get("history", [])
    model = data.get("model") or DEFAULT_MODEL
    file_summaries = data.get("file_summaries", [])

    if not message:
        return jsonify({"error": "Mensaje vacío"}), 400

    messages = []
    for m in history:
        role = m.get("role")
        content = m.get("content")
        if role in ("user", "assistant") and isinstance(content, str):
            messages.append({"role": role, "content": content})

    user_content = message

    if file_summaries:
        user_content += "\n\n[Attached files summary]\n"
        for fsum in file_summaries:
            fname = fsum.get("filename", "file")
            summ = fsum.get("summary", "")
            user_content += f"- {fname}: {summ}\n"

    messages.append({"role": "user", "content": user_content})

    try:
        completion = client.chat.completions.create(
            model=model,
            messages=messages,
        )
        answer = completion.choices[0].message.content
        return jsonify({"reply": answer})
    except Exception as e:
        _log(f"Error in OpenAI chat: {e}")
        return jsonify({"error": "Error calling the AI model."}), 500


@app.route("/upload", methods=["POST"])
def upload():
    """Upload generic files from the main chat and return short summaries."""
    maybe = require_auth()
    if maybe is not None:
        return jsonify({"error": "No autorizado"}), 401

    if "files" not in request.files:
        return jsonify({"error": "No files were sent."}), 400

    files = request.files.getlist("files")
    if not files:
        return jsonify({"error": "No files were sent."}), 400

    results = []
    for f in files:
        filename = secure_filename(f.filename or "file")
        save_name = f"{int(time.time())}_{filename}"
        path = os.path.join(app.config["UPLOAD_FOLDER"], save_name)
        _log(f"Saving uploaded file at {path}")

        f.save(path)

        text = extract_text(path)
        if not text:
            results.append(
                {
                    "filename": filename,
                    "summary": "I could not read this file (unsupported or empty).",
                }
            )
            continue

        try:
            completion = client.chat.completions.create(
                model=DEFAULT_MODEL,
                messages=[
                    {
                        "role": "user",
                        "content": (
                            "Summarize the following document in a few bullet points, "
                            "highlighting key information useful for IT, business analysis "
                            "and project follow-up:\n\n"
                            f"{text[:8000]}"
                        ),
                    }
                ],
            )
            summary = completion.choices[0].message.content
        except Exception as e:
            _log(f"Error summarizing file: {e}")
            summary = (
                "An automatic summary could not be generated, "
                "but the file was uploaded correctly."
            )

        results.append({"filename": filename, "summary": summary})

    return jsonify({"files": results})


# ---------------- DEM / Project Manager API ----------------


def get_dems_filtered(archived: bool):
    dems = load_dems()
    return [enrich_dem(d) for d in dems if bool(d.get("archived", False)) == archived]


@app.route("/api/dems/projects", methods=["GET"])
def list_dems():
    maybe = require_auth()
    if maybe is not None:
        return jsonify({"error": "No autorizado"}), 401

    archived_str = request.args.get("archived", "false").lower()
    archived = archived_str in ("1", "true", "yes")
    projects = get_dems_filtered(archived)
    return jsonify({"projects": projects})


@app.route("/api/dems/projects", methods=["POST"])
def create_dem():
    maybe = require_auth()
    if maybe is not None:
        return jsonify({"error": "No autorizado"}), 401

    data = request.get_json() or {}
    now_iso = datetime.utcnow().isoformat()

    dem = {
        "id": f"dem_{int(time.time() * 1000)}",
        "name": data.get("name", "").strip(),
        "sponsor": data.get("sponsor", "").strip(),
        "requester": data.get("requester", "").strip(),
        "ba_owner": data.get("ba_owner", "").strip(),
        "title": data.get("title", "").strip(),
        "change_request": data.get("change_request", "").strip(),
        "cost_center": data.get("cost_center", "").strip(),
        "status": data.get("status", "").strip() or "Idea",
        "workflow_status": data.get("workflow_status", "").strip() or "Intake",
        "current_owner": data.get("current_owner", "").strip(),
        "start_date": data.get("start_date", "").strip(),
        "priority": (data.get("priority") or "2").strip(),
        "notes": [],
        "documents": [],  # multiple documents per DEM
        "doc_summary": "",
        "created_at": now_iso,
        "updated_at": now_iso,
        "archived": False,
    }

    initial_note = data.get("initial_note", "").strip() if "initial_note" in data else ""
    if initial_note:
        dem["notes"].append(
            {
                "text": _clean_note_text(initial_note),
                "date": datetime.utcnow().strftime("%Y-%m-%d %H:%M"),
            }
        )

    dems = load_dems()
    dems.append(dem)
    save_dems(dems)

    return jsonify({"project": enrich_dem(dem)})


def _update_dem(id, updater):
    dems = load_dems()
    for i, d in enumerate(dems):
        if d.get("id") == id:
            updater(d)
            d["updated_at"] = datetime.utcnow().isoformat()
            dems[i] = d
            save_dems(dems)
            return enrich_dem(d)
    return None


@app.route("/api/dems/projects/<id>/note", methods=["POST"])
def add_dem_note(id):
    maybe = require_auth()
    if maybe is not None:
        return jsonify({"error": "No autorizado"}), 401

    data = request.get_json() or {}
    text = (data.get("text") or "").strip()
    if not text:
        return jsonify({"error": "Nota vacía."}), 400

    clean_text = _clean_note_text(text)

    def updater(d):
        notes = d.get("notes") or []
        if not isinstance(notes, list):
            notes = []
        notes.append(
            {
                "text": clean_text,
                "date": datetime.utcnow().strftime("%Y-%m-%d %H:%M"),
            }
        )
        d["notes"] = notes

    project = _update_dem(id, updater)
    if not project:
        return jsonify({"error": "DEM no encontrado."}), 404
    return jsonify({"project": project})


# ---- EDIT NOTE ---------------------------------------------------------
@app.route("/api/dems/projects/<id>/note/edit", methods=["POST"])
def edit_dem_note(id):
    maybe = require_auth()
    if maybe is not None:
        return jsonify({"error": "Unauthorized"}), 401

    data = request.get_json() or {}
    index = data.get("index")
    new_text = (data.get("text") or "").strip()

    if index is None or new_text == "":
        return jsonify({"error": "Invalid index or empty text."}), 400

    clean_text = _clean_note_text(new_text)

    def updater(d):
        notes = d.get("notes") or []
        if not isinstance(notes, list):
            notes = []
        if index < 0 or index >= len(notes):
            raise ValueError("Invalid index")

        notes[index] = {
            "text": clean_text,
            "date": datetime.utcnow().strftime("%Y-%m-%d %H:%M"),
        }
        d["notes"] = notes

    try:
        project = _update_dem(id, updater)
    except Exception as e:
        return jsonify({"error": str(e)}), 400

    if not project:
        return jsonify({"error": "DEM not found"}), 404
    return jsonify({"project": project})


# ---- DELETE NOTE ---------------------------------------------------------
@app.route("/api/dems/projects/<id>/note/delete", methods=["POST"])
def delete_dem_note(id):
    maybe = require_auth()
    if maybe is not None:
        return jsonify({"error": "Unauthorized"}), 401

    data = request.get_json() or {}
    index = data.get("index")

    if index is None:
        return jsonify({"error": "Missing index"}), 400

    def updater(d):
        notes = d.get("notes") or []
        if not isinstance(notes, list):
            notes = []

        if index < 0 or index >= len(notes):
            raise ValueError("Invalid index")

        notes.pop(index)
        d["notes"] = notes

    try:
        project = _update_dem(id, updater)
    except Exception as e:
        return jsonify({"error": str(e)}), 400

    if not project:
        return jsonify({"error": "DEM not found"}), 404

    return jsonify({"project": project})


@app.route("/api/dems/projects/<id>/update", methods=["POST"])
def update_dem(id):
    maybe = require_auth()
    if maybe is not None:
        return jsonify({"error": "No autorizado"}), 401

    data = request.get_json() or {}

    def updater(d):
        for field in [
            "name",
            "sponsor",
            "requester",
            "ba_owner",
            "title",
            "change_request",
            "cost_center",
            "status",
            "workflow_status",
            "current_owner",
            "start_date",
            "priority",
        ]:
            if field in data:
                d[field] = (data.get(field) or "").strip()

    project = _update_dem(id, updater)
    if not project:
        return jsonify({"error": "DEM no encontrado."}), 404
    return jsonify({"project": project})


@app.route("/api/dems/projects/<id>/archive", methods=["POST"])
def archive_dem(id):
    maybe = require_auth()
    if maybe is not None:
        return jsonify({"error": "No autorizado"}), 401

    def updater(d):
        d["archived"] = True

    project = _update_dem(id, updater)
    if not project:
        return jsonify({"error": "DEM no encontrado."}), 404
    return jsonify({"project": project})


@app.route("/api/dems/projects/<id>/restore", methods=["POST"])
def restore_dem(id):
    maybe = require_auth()
    if maybe is not None:
        return jsonify({"error": "No autorizado"}), 401

    def updater(d):
        d["archived"] = False

    project = _update_dem(id, updater)
    if not project:
        return jsonify({"error": "DEM no encontrado."}), 404
    return jsonify({"project": project})


@app.route("/api/dems/projects/<id>/delete", methods=["POST"])
def delete_dem(id):
    maybe = require_auth()
    if maybe is not None:
        return jsonify({"error": "No autorizado"}), 401

    dems = load_dems()
    new_dems = [d for d in dems if d.get("id") != id]
    if len(new_dems) == len(dems):
        return jsonify({"error": "DEM no encontrado."}), 404
    save_dems(new_dems)
    return jsonify({"success": True})


@app.route("/api/dems/projects/<id>/attach", methods=["POST"])
def attach_doc(id):
    """Attach a document to a DEM, analyze it and store the summary."""
    maybe = require_auth()
    if maybe is not None:
        return jsonify({"error": "No autorizado"}), 401

    if "file" not in request.files:
        return jsonify({"error": "No se recibió archivo."}), 400

    file = request.files["file"]
    filename = secure_filename(file.filename or "document")
    save_name = f"{int(time.time())}_{filename}"
    path = os.path.join(app.config["UPLOAD_FOLDER"], save_name)
    file.save(path)

    text = extract_text(path)
    if not text:
        return jsonify(
            {
                "error": "I could not read this file to generate an executive summary."
            }
        ), 400

    summary = ""
    try:
        completion = client.chat.completions.create(
            model=DEFAULT_MODEL,
            messages=[
                {
                    "role": "user",
                    "content": (
                        "Create an executive summary of the following document for a DEM "
                        "(Digital Enhancement Management) portfolio. "
                        "Focus on: business problem, scope, key requirements, risks, "
                        "dependencies, recommended IT solutions (for example SAP S/4HANA "
                        "or other enterprise systems), and clear next actions. "
                        "Write in concise, professional English. "
                        "Do NOT mention that this text was generated by any AI model and "
                        "do not describe any internal technical process.\n\n"
                        f"{text[:8000]}"
                    ),
                }
            ],
        )
        summary = completion.choices[0].message.content
    except Exception as e:
        _log(f"Error generating doc summary: {e}")
        summary = (
            "An automatic executive summary could not be generated, "
            "but the document was attached to this DEM."
        )

    def updater(d):
        # Ensure documents list exists
        docs = d.get("documents") or []
        if not isinstance(docs, list):
            docs = []
        docs.append(
            {
                "filename": filename,
                "summary": summary,
                "date": datetime.utcnow().strftime("%Y-%m-%d %H:%M"),
            }
        )
        d["documents"] = docs
        # Keep last summary in doc_summary for current UI button
        d["doc_summary"] = summary

    project = _update_dem(id, updater)
    if not project:
        return jsonify({"error": "DEM no encontrado."}), 404
    return jsonify({"project": project})


@app.route("/api/dems/export", methods=["GET"])
def export_active_excel():
    maybe = require_auth()
    if maybe is not None:
        return jsonify({"error": "No autorizado"}), 401

    if Workbook is None:
        return jsonify({"error": "openpyxl no está disponible en el servidor."}), 500

    wb = Workbook()
    ws = wb.active
    ws.title = "Active DEMs"

    headers = [
        "ID",
        "Name",
        "Title",
        "Sponsor",
        "Requester",
        "BA Owner",
        "Cost Center",
        "Status",
        "Workflow Status",
        "Current Task Owner",
        "Start Date",
        "Duration Days",
        "SLA",
        "Last Note",
    ]
    ws.append(headers)

    for dem in get_dems_filtered(False):
        ws.append(
            [
                dem.get("id"),
                dem.get("name"),
                dem.get("title"),
                dem.get("sponsor"),
                dem.get("requester"),
                dem.get("ba_owner"),
                dem.get("cost_center"),
                dem.get("status"),
                dem.get("workflow_status"),
                dem.get("current_owner"),
                dem.get("start_date"),
                dem.get("duration_days"),
                "Breached" if dem.get("sla_breached") else "OK",
                dem.get("last_note"),
            ]
        )

    bio = io.BytesIO()
    wb.save(bio)
    bio.seek(0)
    return send_file(
        bio,
        as_attachment=True,
        download_name="dems_active.xlsx",
        mimetype="application/"
        "vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


@app.route("/api/dems/export_archived", methods=["GET"])
def export_archived_excel():
    maybe = require_auth()
    if maybe is not None:
        return jsonify({"error": "No autorizado"}), 401

    if Workbook is None:
        return jsonify({"error": "openpyxl no está disponible en el servidor."}), 500

    wb = Workbook()
    ws = wb.active
    ws.title = "Archived DEMs"

    headers = [
        "ID",
        "Name",
        "Title",
        "Sponsor",
        "Requester",
        "BA Owner",
        "Cost Center",
        "Status",
        "Workflow Status",
        "Current Task Owner",
        "Start Date",
        "Duration Days",
        "SLA",
        "Last Note",
    ]
    ws.append(headers)

    for dem in get_dems_filtered(True):
        ws.append(
            [
                dem.get("id"),
                dem.get("name"),
                dem.get("title"),
                dem.get("sponsor"),
                dem.get("requester"),
                dem.get("ba_owner"),
                dem.get("cost_center"),
                dem.get("status"),
                dem.get("workflow_status"),
                dem.get("current_owner"),
                dem.get("start_date"),
                dem.get("duration_days"),
                "Breached" if dem.get("sla_breached") else "OK",
                dem.get("last_note"),
            ]
        )

    bio = io.BytesIO()
    wb.save(bio)
    bio.seek(0)
    return send_file(
        bio,
        as_attachment=True,
        download_name="dems_archived.xlsx",
        mimetype="application/"
        "vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


# -------- Export / Import JSON (backup) ----------------------


@app.route("/api/dems/export_json", methods=["GET"])
def export_dems_json():
    """
    Exporta TODOS los DEMs (activos + archivados) como backup JSON.
    El botón “Export JSON” del frontend debe llamar a este endpoint.
    """
    maybe = require_auth()
    if maybe is not None:
        return jsonify({"error": "No autorizado"}), 401

    dems = load_dems() or []
    payload = json.dumps(dems, ensure_ascii=False, indent=2)
    bio = io.BytesIO(payload.encode("utf-8"))
    bio.seek(0)
    return send_file(
        bio,
        as_attachment=True,
        download_name="dems_backup.json",
        mimetype="application/json; charset=utf-8",
    )


@app.route("/api/dems/import", methods=["POST"])
def import_dems_json():
    """
    Importa DEMs desde un JSON.

    El frontend normalmente manda algo así:
        POST /api/dems/import
        { "projects": [ {..dem1..}, {..dem2..}, ... ] }
    """
    maybe = require_auth()
    if maybe is not None:
        return jsonify({"error": "No autorizado"}), 401

    data = request.get_json(silent=True) or {}
    projects = data.get("projects")

    if not isinstance(projects, list):
        return jsonify(
            {"error": "Invalid JSON structure: 'projects' must be a list."}
        ), 400

    current = load_dems() or []
    by_id = {str(d.get("id")): d for d in current if d.get("id")}

    for incoming in projects:
        if not isinstance(incoming, dict):
            continue
        pid = str(incoming.get("id") or "").strip()
        if not pid:
            pid = f"dem_{int(time.time() * 1000)}"
            incoming["id"] = pid
        by_id[pid] = incoming

    merged_list = list(by_id.values())
    save_dems(merged_list)

    enriched = [enrich_dem(d) for d in merged_list]
    return jsonify({"projects": enriched})


# -------- Reporte resumen para panel y descargas -------------


@app.route("/api/dems/report", methods=["POST"])
def dem_report():
    """Texto del reporte para el panel (solo DEMs activos)."""
    maybe = require_auth()
    if maybe is not None:
        return jsonify({"error": "No autorizado"}), 401

    dems = load_dems()
    active_dems = [d for d in dems if not d.get("archived", False)]
    text = build_portfolio_text(active_dems)

    return jsonify({"report": text})


@app.route("/api/dems/download/<fmt>", methods=["GET"])
def dem_download(fmt):
    """Descarga el reporte como TXT / DOCX / PDF (solo DEMs activos)."""
    maybe = require_auth()
    if maybe is not None:
        return jsonify({"error": "No autorizado"}), 401

    fmt = fmt.lower()
    if fmt not in ("txt", "pdf", "docx"):
        return jsonify({"error": "Formato no soportado."}), 400

    dems = [d for d in load_dems() if not d.get("archived", False)]
    text = build_portfolio_text(dems)

    lines = text.split("\n")
    title = lines[0] if lines else "Andres Villanueva DEMS Report"
    body = "\n".join(lines[1:]) if len(lines) > 1 else ""

    if fmt == "txt":
        bio = io.BytesIO()
        bio.write(text.encode("utf-8"))
        bio.seek(0)
        return send_file(
            bio,
            as_attachment=True,
            download_name="dems_portfolio.txt",
            mimetype="text/plain; charset=utf-8",
        )

    if fmt == "docx":
        if Document is None:
            return jsonify({"error": "python-docx no está disponible."}), 500
        doc = Document()
        doc.add_heading(title, level=1)
        doc.add_paragraph("")
        for line in body.split("\n"):
            doc.add_paragraph(line)

        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        return send_file(
            bio,
            as_attachment=True,
            download_name="dems_portfolio.docx",
            mimetype=(
                "application/"
                "vnd.openxmlformats-officedocument.wordprocessingml.document"
            ),
        )

    if fmt == "pdf":
        if SimpleDocTemplate is None:
            return jsonify({"error": "reportlab no está disponible."}), 500

        bio = io.BytesIO()
        doc = SimpleDocTemplate(bio, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        story.append(Paragraph(title, styles["Title"]))
        story.append(Spacer(1, 12))
        for paragraph in body.split("\n\n"):
            story.append(
                Paragraph(paragraph.replace("\n", "<br />"), styles["Normal"])
            )
            story.append(Spacer(1, 8))
        doc.build(story)
        bio.seek(0)
        return send_file(
            bio,
            as_attachment=True,
            download_name="dems_portfolio.pdf",
            mimetype="application/pdf",
        )

    return jsonify({"error": "Formato no soportado."}), 400


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=8080)
